"""Export Word générique — piloté par template + instance (aucun chapitre métier hardcodé)."""

from __future__ import annotations

import copy
import io
import json
from pathlib import Path
from typing import Any

from app.core.database import get_db_path
from app.services.avis_technique_bindings import resolve_calculs, resolve_documents
from app.services.demande_document_storage_service import normalize_stored_path

try:
    from docx import Document
    from docx.shared import Cm, Pt
except ImportError:  # pragma: no cover
    Document = None  # type: ignore
    Cm = None  # type: ignore
    Pt = None  # type: ignore


def _project_root() -> Path:
    # backend/current_fastapi/app/services → repo root
    return Path(__file__).resolve().parents[4]


def _storage_root() -> Path:
    return _project_root() / "storage"


def _resolve_media_path(stored_path: str) -> Path | None:
    raw = normalize_stored_path(stored_path or "")
    if not raw:
        return None
    path = Path(raw)
    if not path.is_absolute():
        path = _storage_root() / raw
        if not path.exists():
            path = _storage_root() / "documents" / raw
    return path if path.exists() and path.is_file() else None


def _add_heading(doc: Any, text: str, level: int = 1) -> None:
    doc.add_heading(str(text or "").strip() or "Section", level=min(max(level, 1), 3))


def _add_paragraph(doc: Any, text: str) -> None:
    p = doc.add_paragraph(str(text or ""))
    if Pt is not None:
        for run in p.runs:
            run.font.size = Pt(10)


def _element_xml(element: Any) -> str:
    try:
        from lxml import etree

        return etree.tostring(element, encoding="unicode")
    except Exception:
        try:
            return str(element.xml)
        except Exception:
            return ""


def _is_cover_sdt(element: Any) -> bool:
    """Word Cover Pages content control (folha de rosto NGE)."""
    xml = _element_xml(element)
    if not xml:
        return False
    return "Cover Pages" in xml or ("docPartGallery" in xml and "Cover" in xml)


def _extract_cover_sdt(body: Any) -> Any | None:
    for child in list(body):
        tag = child.tag.split("}")[-1] if hasattr(child, "tag") else ""
        if tag == "sdt" and _is_cover_sdt(child):
            return copy.deepcopy(child)
    return None


def _iter_w_t(element: Any):
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
    for node in element.iter():
        if getattr(node, "tag", None) == ns:
            yield node


def _fill_cover_sdt(sdt: Any, *, site_title: str, doc_line: str, date_line: str) -> None:
    """Replace the main text runs of the NGE cover SDT (keeps logos / layout)."""
    texts = [n for n in _iter_w_t(sdt) if (n.text or "").strip()]
    if not texts:
        return
    replacements = [site_title, doc_line, date_line]
    idx = 0
    for node in texts:
        raw = (node.text or "").strip()
        if raw.lower().startswith("www."):
            continue
        if idx < len(replacements) and replacements[idx]:
            node.text = replacements[idx]
            idx += 1


def _render_fallback_cover(
    doc: Any,
    *,
    site_title: str,
    doc_line: str,
    date_line: str,
) -> None:
    """Simple folha de rosto when no Cover Pages SDT is available."""
    logo = _project_root() / "frontend" / "react" / "public" / "assets" / "logos" / "nge-logo.png"
    if logo.exists() and Cm is not None:
        try:
            doc.add_picture(str(logo), width=Cm(4.5))
        except Exception:
            pass
    p = doc.add_paragraph()
    run = p.add_run(site_title or "Note technique")
    if Pt is not None:
        run.font.size = Pt(22)
        run.bold = True
    p2 = doc.add_paragraph()
    r2 = p2.add_run(doc_line)
    if Pt is not None:
        r2.font.size = Pt(14)
    if date_line:
        p3 = doc.add_paragraph()
        r3 = p3.add_run(date_line)
        if Pt is not None:
            r3.font.size = Pt(12)
    doc.add_paragraph("www.nge.fr")


def _cover_fields(instance: dict[str, Any], template: dict[str, Any], cover_content: dict[str, Any]) -> dict[str, str]:
    fields = cover_content.get("fields") if isinstance(cover_content.get("fields"), dict) else {}
    meta = instance.get("meta") if isinstance(instance.get("meta"), dict) else {}
    title = str(
        fields.get("title")
        or meta.get("title")
        or instance.get("titre")
        or template.get("label")
        or "Note technique"
    ).strip()
    reference = str(fields.get("reference") or meta.get("reference") or instance.get("reference") or "").strip()
    date_line = str(fields.get("document_date") or meta.get("document_date") or meta.get("date") or "").strip()
    kind = "Note technique"
    label = str(template.get("label") or "").lower()
    if "avis" in label and "nt g3" not in label and "voiries" not in label:
        kind = "Avis technique"
    doc_line = f"{kind} - {reference}" if reference else kind
    return {"site_title": title, "doc_line": doc_line, "date_line": date_line}


def _render_rich_text(doc: Any, content: dict[str, Any]) -> None:
    text = str(content.get("text") or "").strip()
    if text:
        for line in text.splitlines() or [text]:
            _add_paragraph(doc, line)


def _render_bullet_list(doc: Any, content: dict[str, Any]) -> None:
    for item in content.get("items") or []:
        label = item if isinstance(item, str) else str((item or {}).get("text") or item or "")
        if label.strip():
            doc.add_paragraph(label.strip(), style="List Bullet")


def _render_checklist(doc: Any, content: dict[str, Any]) -> None:
    for item in content.get("items") or []:
        if isinstance(item, dict):
            checked = "☑" if item.get("done") or item.get("checked") else "☐"
            label = str(item.get("text") or item.get("label") or "")
            doc.add_paragraph(f"{checked} {label}".strip())
        else:
            doc.add_paragraph(f"☐ {item}")


def _render_kv_table(doc: Any, content: dict[str, Any]) -> None:
    rows = content.get("rows") or []
    if not rows:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Clé"
    hdr[1].text = "Valeur"
    for row in rows:
        if isinstance(row, dict):
            cells = table.add_row().cells
            cells[0].text = str(row.get("key") or "")
            cells[1].text = str(row.get("value") or "")


def _render_one_free_table(doc: Any, *, headers: list[str], rows: list[Any], caption: str = "") -> None:
    cols = len(headers) or (len(rows[0]) if rows and isinstance(rows[0], (list, tuple)) else 0)
    if cols <= 0:
        return
    if caption:
        _add_paragraph(doc, caption)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    if headers:
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
    else:
        for i in range(cols):
            table.rows[0].cells[i].text = f"Col {i + 1}"
    for row in rows:
        cells = table.add_row().cells
        values = row if isinstance(row, (list, tuple)) else [row]
        for i in range(cols):
            cells[i].text = str(values[i] if i < len(values) else "")


def _render_free_table(doc: Any, content: dict[str, Any]) -> None:
    tables = content.get("tables")
    if isinstance(tables, list) and tables:
        for item in tables:
            if not isinstance(item, dict):
                continue
            _render_one_free_table(
                doc,
                headers=[str(h) for h in (item.get("headers") or [])],
                rows=item.get("rows") or [],
                caption=str(item.get("caption") or ""),
            )
        return
    _render_one_free_table(
        doc,
        headers=[str(h) for h in (content.get("headers") or [])],
        rows=content.get("rows") or [],
        caption=str(content.get("caption") or ""),
    )


def _as_person_rows(items: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    if not isinstance(items, list):
        return rows
    for item in items:
        if isinstance(item, dict):
            rows.append([str(item.get("nom") or item.get("name") or ""), str(item.get("service") or "")])
        elif item:
            rows.append([str(item), ""])
    return rows


def _render_meta(doc: Any, content: dict[str, Any], instance: dict[str, Any]) -> None:
    """Page qualité NGE : destinataires, contrôle, historique, métadonnées."""
    content = content if isinstance(content, dict) else {}
    fields = content.get("fields") if isinstance(content.get("fields"), dict) else {}
    meta = instance.get("meta") if isinstance(instance.get("meta"), dict) else {}
    merged = {**meta, **fields}

    title = str(merged.get("title") or instance.get("titre") or "").strip()
    reference = str(merged.get("reference") or instance.get("reference") or "").strip()
    author = str(merged.get("author") or instance.get("auteur") or "").strip()
    status = str(merged.get("status") or instance.get("statut") or "").strip()
    doc_date = str(merged.get("document_date") or merged.get("date") or "").strip()

    _add_heading(doc, "Informations qualité du document", level=1)

    # Destinataires (Pour | Copie)
    pour = _as_person_rows(content.get("destinataires_pour") or merged.get("destinataires_pour"))
    copie = _as_person_rows(content.get("destinataires_copie") or merged.get("destinataires_copie"))
    if not pour and merged.get("destinataire"):
        pour = [[str(merged.get("destinataire")), str(merged.get("destinataire_service") or "")]]
    if pour or copie:
        _add_heading(doc, "Destinataires", level=2)
        max_n = max(len(pour), len(copie), 1)
        dest_rows = []
        for i in range(max_n):
            p = pour[i] if i < len(pour) else ["", ""]
            c = copie[i] if i < len(copie) else ["", ""]
            dest_rows.append(
                [
                    p[0],
                    p[1] if len(p) > 1 else "",
                    c[0],
                    c[1] if len(c) > 1 else "",
                ]
            )
        _render_one_free_table(
            doc,
            headers=["Pour — Nom", "Pour — Service", "Copie — Nom", "Copie — Service"],
            rows=dest_rows,
            caption="",
        )

    # Contrôle final
    controle = content.get("controle_final") if isinstance(content.get("controle_final"), dict) else {}
    if not controle and (author or doc_date):
        controle = {"date": doc_date, "nom": author, "signature": ""}
    if controle:
        _add_heading(doc, "Contrôle final", level=2)
        _render_kv_table(
            doc,
            {
                "rows": [
                    {"key": "Date", "value": controle.get("date") or ""},
                    {"key": "Nom", "value": controle.get("nom") or ""},
                    {"key": "Signature", "value": controle.get("signature") or ""},
                ]
            },
        )

    # Historique des modifications
    historique = content.get("historique") if isinstance(content.get("historique"), list) else []
    if not historique and (author or doc_date or status):
        historique = [
            {
                "version": status or "V0",
                "date": doc_date,
                "redige_par": author,
                "controle": "",
                "modifications": "Émission initiale",
            }
        ]
    if historique:
        _add_heading(doc, "Historique des modifications", level=2)
        hist_rows = []
        for row in historique:
            if not isinstance(row, dict):
                continue
            hist_rows.append(
                [
                    str(row.get("version") or ""),
                    str(row.get("date") or ""),
                    str(row.get("redige_par") or row.get("auteur") or ""),
                    str(row.get("controle") or row.get("controle_externe") or ""),
                    str(row.get("modifications") or ""),
                ]
            )
        _render_one_free_table(
            doc,
            headers=["Version", "Date", "Rédigé par", "Contrôle externe", "Modifications"],
            rows=hist_rows,
        )

    # Autres informations
    _add_heading(doc, "Autres informations", level=2)
    info_rows = [
        {"key": "Auteur", "value": author},
        {"key": "Date de référence", "value": doc_date},
        {"key": "Référence", "value": reference},
        {"key": "Statut document", "value": status},
        {"key": "Titre du document", "value": title},
    ]
    skip = {
        "title",
        "author",
        "document_date",
        "date",
        "status",
        "reference",
        "destinataire",
        "destinataire_service",
        "destinataires_pour",
        "destinataires_copie",
    }
    for key, value in merged.items():
        if key in skip or isinstance(value, (list, dict)):
            continue
        if str(value or "").strip():
            info_rows.append({"key": str(key), "value": value})
    _render_kv_table(doc, {"rows": [r for r in info_rows if str(r.get("value") or "").strip()]})


def _docs_by_id(demande_id: int) -> dict[int, dict[str, Any]]:
    return {int(d["id"]): d for d in resolve_documents(demande_id) if d.get("id") is not None}


def _render_media_cards(doc: Any, content: dict[str, Any], demande_id: int) -> None:
    docs = _docs_by_id(demande_id)
    cards = list(content.get("cards") or [])
    cards.sort(key=lambda c: int((c or {}).get("order") or 0))
    for card in cards:
        if not isinstance(card, dict):
            continue
        caption = str(card.get("caption") or "")
        doc_id = card.get("document_id")
        doc_meta = docs.get(int(doc_id)) if doc_id is not None else None
        path = _resolve_media_path((doc_meta or {}).get("stored_path") or "")
        if path is not None:
            try:
                width = Cm(16) if Cm is not None else None
                if width is not None:
                    doc.add_picture(str(path), width=width)
                else:
                    doc.add_picture(str(path))
            except Exception:
                _add_paragraph(doc, f"[Image non insérable: {path.name}]")
        elif doc_meta:
            _add_paragraph(doc, f"[Document: {doc_meta.get('label') or doc_id}]")
        if caption:
            _add_paragraph(doc, caption)


def _render_document_gallery(doc: Any, content: dict[str, Any], demande_id: int) -> None:
    docs = _docs_by_id(demande_id)
    for item in content.get("items") or []:
        if not isinstance(item, dict):
            continue
        doc_id = item.get("document_id")
        meta = docs.get(int(doc_id)) if doc_id is not None else None
        label = item.get("label") or (meta or {}).get("label") or str(doc_id or "")
        caption = item.get("caption") or ""
        line = label if not caption else f"{label} — {caption}"
        doc.add_paragraph(line, style="List Bullet")


def _render_calculs_table(doc: Any, content: dict[str, Any], demande_id: int) -> None:
    all_calcs = {c["id"]: c for c in resolve_calculs(demande_id)}
    ids = content.get("calcul_ids") or []
    rows = []
    for cid in ids:
        calc = all_calcs.get(int(cid))
        if not calc:
            continue
        rows.append(
            [
                calc.get("reference") or "",
                calc.get("nom_calcul") or "",
                calc.get("statut") or "",
                "oui" if calc.get("a_retenir") else "",
                calc.get("avis") or "",
            ]
        )
    _render_free_table(
        doc,
        {
            "headers": ["Référence", "Nom", "Statut", "Retenu", "Avis"],
            "rows": rows,
        },
    )


def _render_calcul_fiches(doc: Any, content: dict[str, Any], demande_id: int) -> None:
    """Insère les fiches Alizé (PDF annexe) page par page — même rendu que /fiche.pdf."""
    from app.repositories.calculs_repository import CalculsRepository

    ids = [int(x) for x in (content.get("calcul_ids") or []) if str(x).isdigit() or isinstance(x, int)]
    # Fallback only when no explicit binding result was provided (avoid duplicating
    # all impression fiches into empty P1–P6 annexes).
    binding_resolved_empty = bool(content.get("auto_from_binding")) and "calcul_ids" in content
    if not ids and demande_id and not binding_resolved_empty:
        # Fallback : pour_impression / retenus / série CAM1
        all_calcs = resolve_calculs(demande_id, {})
        ids = [int(c["id"]) for c in all_calcs if c.get("pour_impression") or c.get("a_retenir")]
        if not ids:
            ids = [
                int(c["id"])
                for c in all_calcs
                if "CAM1" in str(c.get("nom_calcul") or "").upper().replace(" ", "")
            ]
        if not ids:
            ids = [int(c["id"]) for c in all_calcs if c.get("type_calcul") == "alize"]

    if not ids:
        _add_paragraph(doc, "Aucune fiche de calcul liée à cette annexe.")
        return

    repo = CalculsRepository()
    try:
        import fitz  # PyMuPDF
    except ImportError:
        fitz = None

    for index, cid in enumerate(ids):
        detail = repo.get(cid)
        if not detail:
            continue
        title = detail.nom_calcul or detail.reference or f"Calcul {cid}"
        if index > 0:
            try:
                doc.add_page_break()
            except Exception:
                _add_paragraph(doc, "")
        _add_heading(doc, str(title), level=2)

        pdf_bytes = repo.build_fiche_pdf(cid)
        if pdf_bytes and fitz is not None and Cm is not None:
            try:
                pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
                for page in pdf:
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
                    img_bytes = pix.tobytes("png")
                    stream = io.BytesIO(img_bytes)
                    doc.add_picture(stream, width=Cm(16.5))
                pdf.close()
                continue
            except Exception:
                pass

        # Fallback texte structuré
        alize = detail.alize or {}
        layers = alize.get("layers") or []
        results = alize.get("results") or {}
        platform = alize.get("platform") or {}
        traffic = alize.get("traffic") or {}
        _render_free_table(
            doc,
            {
                "headers": ["Couche", "Épaisseur", "E (MPa)", "ν"],
                "rows": [
                    [
                        layer.get("materiau") or "",
                        "" if layer.get("epaisseur") in (None, "") else f"{layer.get('epaisseur')} cm",
                        "" if layer.get("module") is None else str(layer.get("module")),
                        "" if layer.get("poisson") is None else str(layer.get("poisson")),
                    ]
                    for layer in layers
                ],
            },
        )
        _render_kv_table(
            doc,
            {
                "rows": [
                    {"key": "Plateforme", "value": f"{platform.get('classe') or '—'} / {platform.get('module_pf') or '—'} MPa"},
                    {"key": "NE", "value": str(traffic.get("ne_retenu") or traffic.get("ne_calcule") or results.get("ne") or "—")},
                    {"key": "εt calc / adm", "value": f"{results.get('epsT_calc')} / {results.get('epsT_adm')}"},
                    {"key": "εz calc / adm", "value": f"{results.get('epsZ_calc')} / {results.get('epsZ_adm')}"},
                    {"key": "Conclusion", "value": results.get("conclusion") or ""},
                ]
            },
        )


def _render_materiau_status(doc: Any, content: dict[str, Any]) -> None:
    rows = []
    for item in content.get("items") or []:
        if not isinstance(item, dict):
            continue
        rows.append(
            [
                item.get("materiau") or "",
                item.get("formulation") or "",
                "" if item.get("module") is None else str(item.get("module")),
                item.get("origine") or "",
                item.get("status") or "",
            ]
        )
    _render_free_table(
        doc,
        {
            "headers": ["Matériau", "Formulation", "Module", "Origine", "Statut"],
            "rows": rows,
        },
    )


def render_block(doc: Any, block_type: str, content: dict[str, Any], instance: dict[str, Any]) -> None:
    demande_id = int(instance.get("demande_id") or 0)
    content = content if isinstance(content, dict) else {}
    if block_type == "rich_text":
        _render_rich_text(doc, content)
    elif block_type == "bullet_list":
        _render_bullet_list(doc, content)
    elif block_type == "checklist":
        _render_checklist(doc, content)
    elif block_type == "key_value_table":
        _render_kv_table(doc, content)
    elif block_type == "free_table":
        _render_free_table(doc, content)
    elif block_type == "meta_document":
        _render_meta(doc, content, instance)
    elif block_type == "media_cards":
        _render_media_cards(doc, content, demande_id)
    elif block_type == "document_gallery":
        _render_document_gallery(doc, content, demande_id)
    elif block_type == "calculs_table":
        _render_calculs_table(doc, content, demande_id)
    elif block_type == "calcul_fiches":
        _render_calcul_fiches(doc, content, demande_id)
    elif block_type == "materiau_status":
        _render_materiau_status(doc, content)
    else:
        _add_paragraph(doc, json.dumps(content, ensure_ascii=False)[:2000])


def _slot_is_empty(block_type: str, content: dict[str, Any]) -> bool:
    if block_type == "rich_text":
        return not str(content.get("text") or "").strip()
    if block_type in ("bullet_list", "checklist"):
        return not (content.get("items") or [])
    if block_type == "key_value_table":
        return not (content.get("rows") or [])
    if block_type == "free_table":
        tables = content.get("tables")
        if isinstance(tables, list):
            return not any((t.get("rows") or t.get("headers")) for t in tables if isinstance(t, dict))
        return not (content.get("rows") or content.get("headers"))
    if block_type == "media_cards":
        return not (content.get("cards") or [])
    if block_type == "document_gallery":
        return not (content.get("items") or [])
    if block_type in ("calculs_table", "calcul_fiches"):
        return False  # may resolve from demande
    if block_type == "meta_document":
        fields = content.get("fields") or {}
        return not any(str(v or "").strip() for v in fields.values()) if isinstance(fields, dict) else True
    if block_type == "materiau_status":
        return not (content.get("items") or [])
    return not content


def _is_annex_section(section: dict[str, Any]) -> bool:
    sid = str(section.get("id") or "").lower()
    title = str(section.get("title") or "").lower()
    return sid.startswith("annexe") or sid.startswith("a_") or "annexe" in title


def _page_break(doc: Any) -> None:
    try:
        doc.add_page_break()
    except Exception:
        pass


def _render_sommaire(doc: Any, sections: list[dict[str, Any]]) -> None:
    _add_heading(doc, "Sommaire", level=1)
    for section in sections:
        if not isinstance(section, dict):
            continue
        sid = str(section.get("id") or "")
        if sid == "cover":
            continue
        title = str(section.get("title") or "").strip()
        if not title:
            continue
        # Only top-level chapters + annexes in sommaire (like NT0002 TOC)
        parent = section.get("parent_id")
        if parent is None or _is_annex_section(section):
            doc.add_paragraph(title)


def build_avis_docx_bytes(instance: dict[str, Any], template: dict[str, Any] | None = None) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx n'est pas installé")

    template = template or instance.get("template") or {}
    definition = template.get("definition") if isinstance(template.get("definition"), dict) else {}
    style_path = str(template.get("docx_style_path") or "").strip()
    export_empty = bool(definition.get("export_empty_sections", False))
    export_sommaire = bool(definition.get("export_sommaire", False))

    # Prefer template style donor, else avis optimisation reference
    ref_docx = _project_root() / "storage" / "documents" / (
        "2026-RA-049_Avis_technique_Optimisation_chaussee_RST_D0100_Complet_V2.docx"
    )
    if style_path:
        style_file = Path(style_path)
        if not style_file.is_absolute():
            style_file = _project_root() / style_path
    elif ref_docx.exists():
        style_file = ref_docx
    else:
        style_file = None

    if style_file and style_file.exists():
        doc = Document(str(style_file))
        body = doc.element.body
        # Preserve NGE Cover Pages SDT (folha de rosto) before clearing body
        cover_sdt = _extract_cover_sdt(body)
        for child in list(body):
            if child.tag.endswith("sectPr"):
                continue
            body.remove(child)
    else:
        doc = Document()
        cover_sdt = None

    title = instance.get("titre") or template.get("label") or "Avis technique"
    ref = instance.get("reference") or ""
    doc.core_properties.title = str(title)
    if ref:
        doc.core_properties.subject = str(ref)

    sections = list(definition.get("sections") or [])
    sections.sort(key=lambda s: int((s or {}).get("order") or 0))
    contents = instance.get("contents") if isinstance(instance.get("contents"), dict) else {}
    annex_started = False

    # ── Folha de rosto (Cover Pages) + qualité ────────────────────────────
    cover_slot = (contents.get("cover") or {}).get("meta") if isinstance(contents.get("cover"), dict) else None
    cover_content = (cover_slot or {}).get("content") if isinstance(cover_slot, dict) else {}
    if not isinstance(cover_content, dict):
        cover_content = {}
    if not cover_content.get("fields"):
        cover_content = {
            **cover_content,
            "fields": {
                "title": title,
                "reference": ref,
                "author": instance.get("auteur") or "",
                "status": instance.get("statut") or "",
                **(instance.get("meta") if isinstance(instance.get("meta"), dict) else {}),
            },
        }
    cf = _cover_fields(instance, template, cover_content)
    if cover_sdt is not None:
        _fill_cover_sdt(
            cover_sdt,
            site_title=cf["site_title"],
            doc_line=cf["doc_line"],
            date_line=cf["date_line"],
        )
        body = doc.element.body
        # Insert before final sectPr
        inserted = False
        for i, child in enumerate(list(body)):
            if child.tag.endswith("sectPr"):
                body.insert(i, cover_sdt)
                inserted = True
                break
        if not inserted:
            body.append(cover_sdt)
        _page_break(doc)
    else:
        _render_fallback_cover(
            doc,
            site_title=cf["site_title"],
            doc_line=cf["doc_line"],
            date_line=cf["date_line"],
        )
        _page_break(doc)

    render_block(doc, "meta_document", cover_content, instance)
    _page_break(doc)

    if export_sommaire:
        _render_sommaire(doc, sections)
        _page_break(doc)

    for section in sections:
        if not isinstance(section, dict):
            continue
        section_id = str(section.get("id") or "")
        if section_id == "cover":
            continue

        section_contents = contents.get(section_id) or {}
        blocks = section.get("blocks") or []
        navigable_only = bool(section.get("navigable_only"))
        optional = bool(section.get("optional"))
        is_annex = _is_annex_section(section)

        # Collect non-empty rendered blocks first
        pending: list[tuple[str, str, dict]] = []
        for block_def in blocks:
            if not isinstance(block_def, dict):
                continue
            slot_id = str(block_def.get("slot_id") or "")
            block_type = str(block_def.get("block_type") or "")
            slot = section_contents.get(slot_id) if isinstance(section_contents, dict) else None
            content = (slot or {}).get("content") if isinstance(slot, dict) else {}
            bt = str(((slot or {}).get("block_type") if isinstance(slot, dict) else None) or block_type)
            content = content if isinstance(content, dict) else {}
            if bt not in ("calcul_fiches", "calculs_table") and _slot_is_empty(bt, content):
                continue
            pending.append((str(block_def.get("label") or ""), bt, content))

        if navigable_only and not pending:
            if is_annex and not annex_started:
                _page_break(doc)
                annex_started = True
            _add_heading(doc, str(section.get("title") or section_id), level=1)
            continue

        if not pending and not navigable_only:
            if optional and not export_empty:
                continue
            if not export_empty and not optional:
                # Legacy templates: skip empty body sections
                continue
            # Keep outline (NT G3): emit title + placeholder
            if is_annex and not annex_started:
                _page_break(doc)
                annex_started = True
            elif is_annex:
                _page_break(doc)
            level = 1 if (section.get("parent_id") is None or is_annex) else 2
            _add_heading(doc, str(section.get("title") or section_id or "Section"), level=level)
            _add_paragraph(doc, "À compléter.")
            continue

        if is_annex and not annex_started:
            _page_break(doc)
            annex_started = True
        elif is_annex:
            _page_break(doc)

        level = 1 if (section.get("parent_id") is None or is_annex) else 2
        _add_heading(doc, str(section.get("title") or section_id or "Section"), level=level)

        for label, bt, content in pending:
            if label and label not in {"Texte", "Commentaire", "Introduction", "Fiches de calcul", "Synthèse", "Conclusion"}:
                _add_heading(doc, label, level=min(level + 1, 3))
            render_block(doc, bt, content, instance)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def sanitize_docx_basename(value: str, *, fallback: str = "avis_technique") -> str:
    text = str(value or "").strip() or fallback
    for char in '<>:"/\\|?*\n\r\t':
        text = text.replace(char, "-")
    text = " ".join(text.split()).strip(" .")
    return (text or fallback)[:160]

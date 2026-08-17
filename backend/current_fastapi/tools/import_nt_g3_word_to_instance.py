"""Import NT0002 Word → contents da instance avis RaLab (DB)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from app.models.avis_technique import AvisInstanceUpdateSchema
from app.repositories.avis_technique_repository import AvisTechniqueRepository

INSTANCE_ID = 3
DOCX = (
    Path(__file__).resolve().parents[3]
    / "storage"
    / "documents"
    / "2025-RA-008-D0054-NT0002.docx"
)

CHAPTER_RE = re.compile(r"^\d+$")


def _norm(s: str) -> str:
    s = (s or "").strip().lower()
    s = s.replace("’", "'").replace("`", "'")
    s = re.sub(r"\s+", " ", s)
    return s


def _cell_text(cell) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()).strip()


def _uniq_row_cells(row) -> list[str]:
    out: list[str] = []
    for c in row.cells:
        t = _cell_text(c)
        if not out or out[-1] != t:
            out.append(t)
    return out


def _is_chapter_banner(cells: list[str]) -> bool:
    return len(cells) >= 2 and bool(CHAPTER_RE.match(cells[0].strip())) and bool(cells[1].strip())


def _append_rich(contents: dict, section: str, slot: str, text: str) -> None:
    text = (text or "").strip()
    if not text:
        return
    sec = contents.setdefault(section, {})
    existing = sec.get(slot)
    if isinstance(existing, dict) and existing.get("block_type") == "rich_text":
        prev = str((existing.get("content") or {}).get("text") or "").strip()
        merged = f"{prev}\n\n{text}".strip() if prev else text
        sec[slot] = {"block_type": "rich_text", "content": {"text": merged}}
    else:
        sec[slot] = {"block_type": "rich_text", "content": {"text": text}}


def _set_list(contents: dict, section: str, slot: str, text: str, *, checklist: bool = False) -> None:
    items = []
    for line in (text or "").splitlines():
        line = line.strip().lstrip("•·●◦\uf0b7-–—").strip()
        if len(line) >= 2:
            items.append({"text": line, "done": False} if checklist else line)
    if not items and (text or "").strip():
        items = [{"text": text.strip(), "done": False}] if checklist else [text.strip()]
    bt = "checklist" if checklist else "bullet_list"
    contents.setdefault(section, {})[slot] = {"block_type": bt, "content": {"items": items}}


def _set_free(contents: dict, section: str, slot: str, free: dict) -> None:
    contents.setdefault(section, {})[slot] = {"block_type": "free_table", "content": free}


def _split_label_body(cells: list[str]) -> tuple[str, str]:
    if len(cells) >= 2 and cells[0].strip() and (cells[1].strip() or len(cells[0]) < 80):
        return cells[0].strip(), cells[1].strip()
    raw = cells[0] if cells else ""
    if "\n" in raw:
        a, b = raw.split("\n", 1)
        return a.strip(), b.strip()
    return raw.strip(), ""


def _chapter_key(title: str) -> str:
    t = (title or "").upper()
    if "PRÉAMBULE" in t or "PREAMBULE" in t:
        return "ch1"
    if "MARCH" in t:
        return "ch2"
    if "PROPOS" in t and "ANALYSE" not in t:
        return "ch3"
    if "RECONNAISSANCE" in t:
        return "ch4"
    if "ARASE" in t:
        return "ch5"
    if "HYPOTH" in t:
        return "ch6"
    if "VÉRIFICATION" in t or "VERIFICATION" in t:
        return "ch7"
    if "ANALYSE" in t:
        return "ch8"
    if "DISPOSITION" in t:
        return "ch9"
    if "POINT" in t or "ARRÊT" in t or "ARRET" in t:
        return "ch10"
    if "SYNTH" in t:
        return "ch11"
    if "ANNEXE" in t:
        return "annexe"
    return ""


def _route(chapter: str, label: str) -> tuple[str, str, str] | None:
    """Return (section_id, slot_id, mode) mode=rich|list|check."""
    n = _norm(label)
    ch = chapter

    # Explicit by chapter + label
    routes = {
        ("ch1", "présentation"): ("s1_1", "texte", "rich"),
        ("ch1", "cadre de la mission"): ("s1_1", "texte", "rich"),
        ("ch1", "documents pris en compte"): ("s1_2", "texte", "rich"),
        ("ch2", "document marché"): ("s2_1", "texte", "rich"),
        ("ch2", "structures marché"): ("s2_2", "texte", "rich"),
        ("ch3", "structures proposées et adoptées"): ("s3_1", "texte", "rich"),
        ("ch3", "lecture immédiate"): ("s3_1", "texte", "rich"),
        ("ch4", "identification gtr"): ("s4_1", "texte", "rich"),
        ("ch4", "lecture gtr"): ("s4_2", "texte", "rich"),
        ("ch5", "synthèse géotechnique"): ("s5_1", "texte", "rich"),
        ("ch5", "principe g3"): ("s5_2", "texte", "rich"),
        ("ch5", "point de vigilance"): ("s5_2", "texte", "rich"),
        ("ch6", "données retenues"): ("s6_1", "texte", "rich"),
        ("ch6", "entrées de calcul"): ("s6_2", "texte", "rich"),
        ("ch6", "principe de modélisation"): ("s6_2", "texte", "rich"),
        ("ch7", "calage sur cas de référence"): ("s7_1", "texte", "rich"),
        ("ch7", "comparaison"): ("s7_2", "texte", "rich"),
        ("ch7", "lecture du calcul"): ("s7_2", "texte", "rich"),
        ("ch8", "lecture technique"): ("s8_1", "texte", "rich"),
        ("ch8", "structures proposées et adoptées"): ("s8_2", "texte", "rich"),
        ("ch8", "conditions"): ("s8_3", "texte", "rich"),
        ("ch9", "arase"): ("s9_1", "texte", "rich"),
        ("ch9", "couches et réception"): ("s9_2", "texte", "rich"),
        ("ch9", "prescriptions retenues"): ("s9_2", "texte", "rich"),
        ("ch10", "points d'arrêt"): ("s10_1", "texte", "rich"),
        ("ch10", "documents g3"): ("s10_2", "texte", "rich"),
        ("ch11", "structures proposées et adoptées"): ("s11_1", "texte", "rich"),
        ("ch11", "conclusion"): ("s11_2", "texte", "rich"),
        ("annexe", "trafic"): ("annexe_a", "texte", "rich"),
        ("annexe", "résultats"): ("annexe_b", "texte", "rich"),
        ("annexe", "resultats"): ("annexe_b", "texte", "rich"),
    }
    if (ch, n) in routes:
        return routes[(ch, n)]

    # Fuzzy label only
    fuzzy = {
        "présentation": ("s1_1", "texte", "rich"),
        "cadre de la mission": ("s1_1", "texte", "rich"),
        "documents pris en compte": ("s1_2", "texte", "rich"),
        "document marché": ("s2_1", "texte", "rich"),
        "structures marché": ("s2_2", "texte", "rich"),
        "lecture immédiate": ("s3_1", "texte", "rich"),
        "identification gtr": ("s4_1", "texte", "rich"),
        "lecture gtr": ("s4_2", "texte", "rich"),
        "synthèse géotechnique": ("s5_1", "texte", "rich"),
        "principe g3": ("s5_2", "texte", "rich"),
        "point de vigilance": ("s5_2", "texte", "rich"),
        "données retenues": ("s6_1", "texte", "rich"),
        "entrées de calcul": ("s6_2", "texte", "rich"),
        "principe de modélisation": ("s6_2", "texte", "rich"),
        "calage sur cas de référence": ("s7_1", "texte", "rich"),
        "lecture du calcul": ("s7_2", "texte", "rich"),
        "lecture technique": ("s8_1", "texte", "rich"),
        "conditions": ("s8_3", "texte", "rich"),
        "arase": ("s9_1", "texte", "rich"),
        "couches et réception": ("s9_2", "texte", "rich"),
        "prescriptions retenues": ("s9_2", "texte", "rich"),
        "points d'arrêt": ("s10_1", "texte", "rich"),
        "documents g3": ("s10_2", "texte", "rich"),
        "conclusion": ("s11_2", "texte", "rich"),
        "trafic": ("annexe_a", "texte", "rich"),
        "résultats": ("annexe_b", "texte", "rich"),
    }
    if n in fuzzy:
        return fuzzy[n]
    if "structures proposées" in n:
        if ch == "ch8":
            return ("s8_2", "texte", "rich")
        if ch == "ch11":
            return ("s11_1", "texte", "rich")
        return ("s3_1", "texte", "rich")
    if n == "comparaison":
        return ("s7_2", "texte", "rich")
    return None


def _table_as_free(tbl) -> dict | None:
    if len(tbl.rows) < 2 or len(tbl.columns) < 2:
        return None
    first = _uniq_row_cells(tbl.rows[0])
    if _is_chapter_banner(first):
        return None
    joined = " ".join(first).lower()
    if any(x in joined for x in ("pour", "copie", "version", "auteur", "contrôle final", "controle final")):
        return None
    if len(tbl.rows) == 1:
        return None
    headers = first
    rows = []
    for row in tbl.rows[1:]:
        cells = _uniq_row_cells(row)
        if not any(c.strip() for c in cells):
            continue
        while len(cells) < len(headers):
            cells.append("")
        rows.append(cells[: len(headers)])
    if not rows:
        return None
    return {"headers": headers, "rows": rows}


def parse_nt0002(docx_path: Path) -> dict:
    d = Document(str(docx_path))
    contents: dict = {}
    cover_fields: dict = {}
    destinataires_pour: list[dict] = []
    destinataires_copie: list[dict] = []
    historique: list[dict] = []
    controle_final: dict = {}
    chapter = ""

    for tbl in d.tables:
        cells0 = _uniq_row_cells(tbl.rows[0]) if tbl.rows else []

        if cells0 and cells0[0].strip().lower().startswith("pour"):
            for row in tbl.rows[2:]:
                cells = _uniq_row_cells(row)
                flat = list(cells)
                if len(flat) >= 4:
                    if flat[0] or flat[1]:
                        destinataires_pour.append({"nom": flat[0], "service": flat[1]})
                    if flat[2] or flat[3]:
                        destinataires_copie.append({"nom": flat[2], "service": flat[3]})
            continue

        if cells0 and "contrôle final" in _norm(cells0[0]):
            for row in tbl.rows[1:]:
                cells = _uniq_row_cells(row)
                if len(cells) >= 2:
                    key = _norm(cells[0])
                    controle_final[{"date": "date", "nom": "nom", "signature": "signature"}.get(key, key)] = cells[1]
            continue

        if cells0 and cells0[0].strip().lower() == "version" and "date" in _norm(" ".join(cells0)):
            for row in tbl.rows[1:]:
                cells = _uniq_row_cells(row)
                if len(cells) >= 5 and any(cells):
                    historique.append(
                        {
                            "version": cells[0],
                            "date": cells[1],
                            "redige_par": cells[2],
                            "controle": cells[3],
                            "modifications": cells[4],
                        }
                    )
            continue

        if cells0 and _norm(cells0[0]) == "auteur" and len(tbl.rows) >= 2:
            for row in tbl.rows:
                cells = _uniq_row_cells(row)
                if len(cells) < 2:
                    continue
                k, v = _norm(cells[0]), cells[1]
                if "auteur" in k:
                    cover_fields["author"] = v
                elif "date" in k:
                    cover_fields["document_date"] = v
                elif "r" in k and "f" in k:
                    cover_fields["reference"] = v
                elif "statut" in k:
                    cover_fields["status"] = v
                elif "titre" in k:
                    cover_fields["title"] = v
            continue

        if _is_chapter_banner(cells0):
            chapter = _chapter_key(cells0[1])
            continue

        free = _table_as_free(tbl)
        if free is not None:
            targets = {
                "ch1": ("s1_2", "tableau"),
                "ch2": ("s2_2", "tableau"),
                "ch3": ("s3_1", "tableau"),
                "ch4": ("s4_1", "tableau"),
                "ch5": ("s5_1", "proctor"),
                "ch6": ("s6_1", "tableau"),
                "ch7": ("s7_2", "resultats"),
                "ch8": ("s8_2", "tableau"),
                "ch10": ("s10_1", "pa"),
                "annexe": ("annexe_a", "tableau"),
            }
            if chapter in targets:
                sec, slot = targets[chapter]
                _set_free(contents, sec, slot, free)
            continue

        if len(tbl.rows) != 1:
            continue

        label, body = _split_label_body(cells0)
        if not label and not body:
            continue
        if CHAPTER_RE.match(label.strip()) and body:
            chapter = _chapter_key(body)
            continue

        route = _route(chapter, label)
        text = f"{label}\n\n{body}".strip() if label and body else (body or label)
        if not route:
            # chapter fallback
            fallback = {
                "ch1": ("s1_1", "texte", "rich"),
                "ch2": ("s2_1", "texte", "rich"),
                "ch3": ("s3_1", "texte", "rich"),
                "ch4": ("s4_1", "texte", "rich"),
                "ch5": ("s5_1", "texte", "rich"),
                "ch6": ("s6_1", "texte", "rich"),
                "ch7": ("s7_2", "texte", "rich"),
                "ch8": ("s8_2", "texte", "rich"),
                "ch9": ("s9_1", "texte", "rich"),
                "ch10": ("s10_1", "texte", "rich"),
                "ch11": ("s11_1", "texte", "rich"),
            }
            route = fallback.get(chapter)
        if not route:
            continue
        sec, slot, mode = route
        if mode == "list":
            _set_list(contents, sec, slot, body or text)
        elif mode == "check":
            _set_list(contents, sec, slot, body or text, checklist=True)
        else:
            _append_rich(contents, sec, slot, text)
            # Also feed bullet slots for list-looking bodies
            if body and (body.count("\n") >= 2 or "•" in body or "\uf0b7" in body):
                if sec == "s1_1" and slot == "texte":
                    pass
                elif sec == "s5_2":
                    _set_list(contents, sec, "liste", body)
                elif sec == "s8_1":
                    _set_list(contents, sec, "liste", body)
                elif sec == "s8_3":
                    _set_list(contents, sec, "liste", body)
                elif sec == "s9_1":
                    _set_list(contents, sec, "liste", body, checklist=True)
                elif sec == "s9_2":
                    _set_list(contents, sec, "liste", body, checklist=True)
                elif sec == "s10_2":
                    _set_list(contents, sec, "liste", body)

    contents["cover"] = {
        "meta": {
            "block_type": "meta_document",
            "content": {
                "fields": cover_fields,
                "destinataires_pour": destinataires_pour,
                "destinataires_copie": destinataires_copie,
                "controle_final": controle_final,
                "historique": historique,
            },
        }
    }
    return contents


def main() -> None:
    repo = AvisTechniqueRepository()
    inst = repo.get_instance(INSTANCE_ID, with_template=True)
    if not inst:
        raise SystemExit("instance 3 missing")

    imported = parse_nt0002(DOCX)
    existing = dict(inst.get("contents") or {})

    # Reset textual slots then merge import (keep binding blocks)
    keep_types = {"calculs_table", "calcul_fiches", "document_gallery", "materiau_status", "media_cards"}
    for sec_id, slots in list(existing.items()):
        if sec_id in {"cover"}:
            continue
        for slot_id, block in list((slots or {}).items()):
            bt = (block or {}).get("block_type")
            if bt not in keep_types:
                # will be replaced if imported
                pass

    for sec_id, slots in imported.items():
        cur = existing.setdefault(sec_id, {})
        for slot_id, block in (slots or {}).items():
            bt = block.get("block_type")
            if bt in keep_types and slot_id in cur:
                continue
            cur[slot_id] = block

    fields = (((imported.get("cover") or {}).get("meta") or {}).get("content") or {}).get("fields") or {}
    repo.update_instance(
        INSTANCE_ID,
        AvisInstanceUpdateSchema(
            titre=fields.get("title") or inst.get("titre"),
            auteur=fields.get("author") or inst.get("auteur") or "COSTA PEREIRA Marco",
            meta={**(inst.get("meta") or {}), **fields},
            contents=existing,
        ),
    )
    repo.refresh_bindings(INSTANCE_ID, only_empty=True)

    inst = repo.get_instance(INSTANCE_ID, with_template=True)
    contents = inst.get("contents") or {}
    lines = []
    for sid in [
        "cover",
        "s1_1",
        "s1_2",
        "s2_1",
        "s2_2",
        "s3_1",
        "s4_1",
        "s4_2",
        "s5_1",
        "s5_2",
        "s6_1",
        "s6_2",
        "s7_1",
        "s7_2",
        "s8_1",
        "s8_2",
        "s8_3",
        "s9_1",
        "s9_2",
        "s10_1",
        "s10_2",
        "s11_1",
        "s11_2",
        "annexe_a",
        "annexe_b",
    ]:
        slots = contents.get(sid) or {}
        bits = []
        for k, v in slots.items():
            c = (v or {}).get("content") or {}
            bt = (v or {}).get("block_type")
            ok = False
            nchars = 0
            if bt == "rich_text":
                t = str(c.get("text") or "")
                ok = bool(t.strip())
                nchars = len(t.strip())
            elif bt in ("bullet_list", "checklist"):
                ok = bool(c.get("items"))
                nchars = len(c.get("items") or [])
            elif bt == "free_table":
                ok = bool(c.get("rows") or c.get("headers"))
            elif bt == "meta_document":
                ok = bool(c.get("fields") or c.get("destinataires_pour"))
            elif bt in keep_types:
                ok = bool(c.get("items") or c.get("calcul_ids") or c.get("cards"))
            bits.append(f"{k}:{'Y' if ok else 'n'}({nchars})")
        lines.append(f"{sid}: {', '.join(bits) or 'EMPTY'}")

    report = Path("tools/_tmp_import_nt0002_report.txt")
    report.write_text("\n".join(lines), encoding="utf-8")
    print(report.read_text(encoding="utf-8"))
    # samples
    for sid in ("s1_1", "s3_1", "s8_2", "s11_2"):
        t = (((contents.get(sid) or {}).get("texte") or {}).get("content") or {}).get("text") or ""
        Path(f"tools/_tmp_sample_{sid}.txt").write_text(t[:1200], encoding="utf-8")
    print("OK /avis-technique/" + str(INSTANCE_ID))


if __name__ == "__main__":
    main()

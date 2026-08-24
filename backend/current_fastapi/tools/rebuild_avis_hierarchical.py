"""Construit template avis hiérarchique + importe le Word (textes, tableaux, images) en DB."""
from __future__ import annotations

import json
import re
import zipfile
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[3]
DOCX = ROOT / "storage" / "documents" / "2026-RA-049_Avis_technique_Optimisation_chaussee_RST_D0100_Complet_V2.docx"
SEED = Path(__file__).resolve().parents[1] / "app" / "data" / "avis_templates" / "optimisation_chaussee_rst.json"
INSTANCE_ID = 1
AFFAIRE_REF = "2026-RA-049"

# Sections qui portent typiquement des figures dans le Word de référence
SECTIONS_WITH_FIGURES = {
    "s1_3", "s2_1", "s2_2", "s2_3", "s2_4",
    "s3_1", "s3_2", "s3_3", "s3_5",
    "s4_1", "s4_5",
    "a1", "a2", "a3",
}

OUTLINE = [
    {"id": "cover", "title": "Informations qualité du document", "parent_id": None, "order": 0, "kind": "cover"},
    {"id": "ch1", "title": "1. Préambule", "parent_id": None, "order": 10},
    {"id": "s1_1", "title": "1.1. Présentation", "parent_id": "ch1", "order": 11, "match": ["1.1. Présentation", "1.1 Présentation"]},
    {"id": "s1_2", "title": "1.2. Objectifs recherchés", "parent_id": "ch1", "order": 12, "match": ["1.2. Objectifs", "1.2 Objectifs"]},
    {"id": "s1_3", "title": "1.3. Contraintes du site", "parent_id": "ch1", "order": 13, "match": ["1.3. Contraintes", "1.3 Contraintes"]},
    {"id": "s1_4", "title": "1.4. Documents de référence", "parent_id": "ch1", "order": 14, "match": ["1.4. Documents", "1.4 Documents"]},
    {"id": "s1_5", "title": "1.5. Méthodologie appliquée", "parent_id": "ch1", "order": 15, "match": ["1.5. Méthodologie", "1.5 Méthodologie"]},
    {"id": "ch2", "title": "2. Informations sur le projet", "parent_id": None, "order": 20},
    {"id": "s2_1", "title": "2.1. Implantation du projet", "parent_id": "ch2", "order": 21, "match": ["2.1. Implantation", "2.1 Implantation"]},
    {"id": "s2_2", "title": "2.2. Géométrie du projet", "parent_id": "ch2", "order": 22, "match": ["2.2. Géométrie", "2.2 Géométrie"]},
    {"id": "s2_3", "title": "2.3. Environnement", "parent_id": "ch2", "order": 23, "match": ["2.3. Environnement", "2.3 Environnement"]},
    {"id": "s2_4", "title": "2.4. Historique et interfaces", "parent_id": "ch2", "order": 24, "match": ["2.4. Historique", "2.4 Historique"]},
    {"id": "s2_5", "title": "2.5. Contexte géologique et hydrogéologique", "parent_id": "ch2", "order": 25, "match": ["2.5. Contexte géologique", "2.5. Contexte", "2.5 Contexte"]},
    {"id": "ch3", "title": "3. Analyse des données disponibles", "parent_id": None, "order": 30},
    {"id": "s3_1", "title": "3.1. Données du CCTP", "parent_id": "ch3", "order": 31, "match": ["3.1. Données du CCTP", "3.1 Données du CCTP"]},
    {"id": "s3_2", "title": "3.2. Données de l’étude géotechnique", "parent_id": "ch3", "order": 32, "match": ["3.2. Données", "3.2 Données"]},
    {"id": "s3_3", "title": "3.3. Voirie provisoire AFTRAL", "parent_id": "ch3", "order": 33, "match": ["3.3. Voirie", "3.3 Voirie"]},
    {"id": "s3_4", "title": "3.4. Limites des données disponibles", "parent_id": "ch3", "order": 34, "match": ["3.4. Limites", "3.4 Limites"]},
    {"id": "s3_5", "title": "3.5. Conséquences sur la structure provisoire", "parent_id": "ch3", "order": 35, "match": ["3.5. Conséquences", "3.5 Conséquences"]},
    {"id": "ch4", "title": "4. Dimensionnement", "parent_id": None, "order": 40},
    {"id": "s4_1", "title": "4.1. Partie supérieure des terrassements", "parent_id": "ch4", "order": 41, "match": ["4.1. Partie", "4.1 Partie"]},
    {"id": "s4_2", "title": "4.2. Amélioration de la PST et couche de forme", "parent_id": "ch4", "order": 42, "match": ["4.2. Amélioration", "4.2 Amélioration", "4.3."]},
    {"id": "s4_4", "title": "4.4. Hypothèses de trafic", "parent_id": "ch4", "order": 44, "match": ["4.4. Hypothèses", "4.4 Hypothèses"], "prefer_table": True},
    {"id": "s4_5", "title": "4.5. Structure DCE", "parent_id": "ch4", "order": 45, "match": ["4.5. Structure", "4.5 Structure"]},
    {"id": "s4_6", "title": "4.6. Matériaux des variantes BBME et BBSG", "parent_id": "ch4", "order": 46, "match": ["4.6. Matériaux", "4.6 Matériaux"]},
    {"id": "s4_7", "title": "4.7. Variantes BBME et BBSG par rapport au DCE", "parent_id": "ch4", "order": 47, "match": ["4.7. Variantes", "4.7 Variantes"]},
    {"id": "s4_8", "title": "4.8. Vérification mécanique", "parent_id": "ch4", "order": 48, "match": ["4.8. Vérification", "4.8 Vérification"]},
    {"id": "s4_9", "title": "4.9. Résultats des variantes BBME et BBSG", "parent_id": "ch4", "order": 49, "match": ["4.9. Résultats", "4.9 Résultats"]},
    {"id": "s4_10", "title": "4.10. Comparaison et choix de la variante", "parent_id": "ch4", "order": 50, "match": ["4.10. Comparaison", "4.10 Comparaison"]},
    {"id": "ch5", "title": "5. Synthèse et dispositions constructives", "parent_id": None, "order": 60},
    {"id": "s5_1", "title": "5.1. Structures de chaussée proposées", "parent_id": "ch5", "order": 61, "match": ["5.1. Structures", "5.1 Structures"]},
    {"id": "s5_2", "title": "5.2. PST, couche de forme et réception", "parent_id": "ch5", "order": 62, "match": ["5.2. PST", "5.2 PST", "5.3.", "5.4."]},
    {"id": "s5_5", "title": "5.5. Mise en œuvre GB4 et couches de roulement", "parent_id": "ch5", "order": 65, "match": ["5.5. Mise", "5.5 Mise", "5.6."]},
    {"id": "s5_7", "title": "5.7. Contrôles et limites de la proposition", "parent_id": "ch5", "order": 67, "match": ["5.7. Contrôles", "5.7 Contrôles", "5.8."]},
    {"id": "ch6", "title": "6. Annexes", "parent_id": None, "order": 70},
    {"id": "a1", "title": "Annexe 1 - Plan d’aménagement", "parent_id": "ch6", "order": 71, "match": ["Annexe 1"], "prefer_media": True},
    {"id": "a2", "title": "Annexe 2 - Voirie provisoire AFTRAL", "parent_id": "ch6", "order": 72, "match": ["Annexe 2"], "prefer_media": True},
    {"id": "a3", "title": "Annexe 3 - Références matériaux et données de calcul", "parent_id": "ch6", "order": 73, "match": ["Annexe 3"]},
    {"id": "a4", "title": "Annexe 4 - Informations restant à compléter", "parent_id": "ch6", "order": 74, "match": ["Annexe 4"]},
    {"id": "a5", "title": "Annexe 5 - Tableau comparatif final", "parent_id": "ch6", "order": 75, "match": ["Annexe 5"], "prefer_table": True},
    {"id": "a6", "title": "Annexe 6 - Calculs mécaniques complets", "parent_id": "ch6", "order": 76, "match": ["Annexe 6", "ANNEXE 6"], "prefer_calculs": True},
]

SECTION_MATCHERS: list[tuple[str, list[str]]] = [
    (n["id"], n["match"]) for n in OUTLINE if n.get("match")
]


def cell_text(cell) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def unique_row_cells(row) -> list[str]:
    seen: set[int] = set()
    out: list[str] = []
    for cell in row.cells:
        tid = id(cell._tc)
        if tid in seen:
            continue
        seen.add(tid)
        out.append(cell_text(cell))
    return out


def table_matrix(table: Table) -> list[list[str]]:
    return [unique_row_cells(row) for row in table.rows]


def nested_tables(table: Table) -> list[Table]:
    """Tableaux Word imbriqués dans les cellules (ex. 1.4 Documents de référence)."""
    seen: set[int] = set()
    out: list[Table] = []
    for row in table.rows:
        for cell in row.cells:
            for nt in cell.tables:
                tid = id(nt._tbl)
                if tid in seen:
                    continue
                seen.add(tid)
                out.append(nt)
    return out

def is_banner_table(rows: list[list[str]]) -> bool:
    """Chapitre Word en bandeau 1x2 : ['1','PRÉAMBULE'] / ['4','DIMENSIONNEMENT']."""
    nr = len(rows)
    nc = max((len(r) for r in rows), default=0)
    texts = [x for row in rows for x in row if x]
    if not texts:
        return True
    if nr == 1 and nc <= 2 and texts[0].isdigit() and len(texts[0]) <= 2:
        # ex. 1 | PRÉAMBULE  /  6 | ANNEXES
        return True
    if nr <= 2 and nc <= 2:
        joined = " ".join(texts).upper()
        # Uniquement les bandeaux de chapitre (pas « Synthèse géotechnique »)
        if texts[0].isdigit() and any(
            x in joined
            for x in (
                "PRÉAMBULE",
                "INFORMATIONS SUR LE PROJET",
                "ANALYSE DES DONNÉES",
                "DIMENSIONNEMENT",
                "SYNTHÈSE ET DISPOSITIONS",
                "ANNEXES",
                "ANNEXE 6",
            )
        ):
            return True
    return False


def is_toc_line(text: str) -> bool:
    t = str(text or "").strip()
    if not t:
        return True
    if t.lower() == "sommaire":
        return True
    # "1. Préambule\t4" ou "…........22"
    if re.search(r"\t\d+\s*$", t):
        return True
    if re.search(r"\.{4,}\s*\d+\s*$", t):
        return True
    # lignes sommaire courtes type "1.1 Présentation 4"
    if re.match(r"^\d+(\.\d+)*\.?\s+.+\s+\d{1,3}$", t) and len(t) < 80:
        return True
    return False


def is_data_table(rows: list[list[str]]) -> bool:
    if is_banner_table(rows):
        return False
    nr = len(rows)
    nc = max((len(r) for r in rows), default=0)
    nonempty = sum(1 for row in rows for c in row if c)
    if nr < 2 or nc < 2 or nonempty < 3:
        return False
    # Petites grilles de libellés (ex. annexe 2) : pas un vrai tableau de données
    cells = [c for row in rows for c in row if c]
    if nr <= 3 and nc <= 3 and cells and all(len(c) < 60 for c in cells):
        return False
    return True


def detect_section(text: str, current: str) -> str:
    """Match le needle le plus spécifique (plus long) pour éviter 2.6. vs 2.5."""
    t = text.strip()
    best_sid = None
    best_len = -1
    for sid, needles in SECTION_MATCHERS:
        for n in needles:
            if t.startswith(n) or (len(n) >= 6 and n in t[:100]):
                if len(n) > best_len:
                    best_len = len(n)
                    best_sid = sid
    return best_sid or current


def blip_embeds(element) -> list[str]:
    embeds = []
    for blip in element.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid:
            embeds.append(rid)
    return embeds


def resolve_image_parts(doc: Document) -> dict[str, tuple[str, bytes]]:
    """rId -> (filename, bytes) for images in the main document part."""
    out: dict[str, tuple[str, bytes]] = {}
    for rel in doc.part.rels.values():
        reltype = str(getattr(rel, "reltype", "") or "")
        if "image" not in reltype.lower():
            continue
        try:
            blob = rel.target_part.blob
            name = Path(str(rel.target_ref)).name
            out[rel.rId] = (name, blob)
        except Exception:
            continue
    return out


def free_table_from_rows(rows: list[list[str]], caption: str = "") -> dict:
    if not rows:
        return {"caption": caption, "headers": [], "rows": []}
    headers = [str(c or "") for c in rows[0]]
    body = []
    for row in rows[1:]:
        body.append([str(c or "") for c in row])
    # pad
    ncols = len(headers)
    body = [r + [""] * (ncols - len(r)) if len(r) < ncols else r[:ncols] for r in body]
    return {"caption": caption, "headers": headers, "rows": body}


def build_template_definition() -> dict:
    sections = []
    for node in OUTLINE:
        blocks = []
        sid = node["id"]
        if sid == "cover":
            blocks = [{"slot_id": "meta", "block_type": "meta_document", "label": "Qualité document"}]
        elif node.get("parent_id") is None and sid != "cover":
            blocks = []
        elif node.get("prefer_calculs"):
            blocks = [
                {"slot_id": "texte", "block_type": "rich_text", "label": "Introduction"},
                {
                    "slot_id": "fiches",
                    "block_type": "calcul_fiches",
                    "label": "Fiches de calcul",
                    "binding": {"source": "calculs", "filter": {"type_calcul": "alize", "nom_contains": "CAM1"}},
                },
            ]
        elif node.get("prefer_media"):
            blocks = [
                {"slot_id": "texte", "block_type": "rich_text", "label": "Commentaire"},
                {"slot_id": "figures", "block_type": "media_cards", "label": "Plans / photos"},
            ]
        elif node.get("prefer_table"):
            blocks = [
                {"slot_id": "texte", "block_type": "rich_text", "label": "Commentaire"},
                {"slot_id": "tableau", "block_type": "free_table", "label": "Tableau"},
                {"slot_id": "liste", "block_type": "bullet_list", "label": "Points"},
            ]
        else:
            blocks = [
                {"slot_id": "texte", "block_type": "rich_text", "label": "Texte"},
                {"slot_id": "liste", "block_type": "bullet_list", "label": "Liste"},
                {"slot_id": "tableau", "block_type": "free_table", "label": "Tableau"},
            ]
        if sid in SECTIONS_WITH_FIGURES and not any(b.get("slot_id") == "figures" for b in blocks):
            blocks.append({"slot_id": "figures", "block_type": "media_cards", "label": "Figures"})

        sections.append(
            {
                "id": sid,
                "title": node["title"],
                "order": node["order"],
                "parent_id": node.get("parent_id"),
                "optional": False,
                "navigable_only": bool(node.get("parent_id") is None and sid != "cover"),
                "blocks": blocks,
            }
        )
    return {
        "reference_rule": "{annee}-{region}-NT{seq:04d}",
        "meta_field_keys": ["title", "author", "document_date", "status", "destinataire"],
        "sections": sections,
    }


def walk_document(docx_path: Path) -> tuple[dict, dict[str, list[str]], dict[str, tuple[str, bytes]]]:
    """
    Returns:
      buckets[section_id] = {texts, bullets, tables: [free_table dicts], captions_before}
      images_by_section[section_id] = [rId, ...]
      image_parts rId -> (filename, bytes)
    """
    doc = Document(str(docx_path))
    image_parts = resolve_image_parts(doc)
    buckets: dict[str, dict] = defaultdict(lambda: {"texts": [], "bullets": [], "tables": [], "prev_caption": ""})
    images_by_section: dict[str, list[str]] = defaultdict(list)
    current = "cover"
    prev_text = ""
    in_toc = False
    body_started = False

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            text = p.text.strip()
            embeds = blip_embeds(child)
            if embeds and body_started and not in_toc:
                for rid in embeds:
                    images_by_section[current].append(rid)
            if not text:
                continue
            if text.lower() == "sommaire":
                in_toc = True
                continue
            if in_toc:
                # fin du sommaire = première vraie section corps
                if text.startswith("1.1") or text.upper().startswith("PRÉAMBULE"):
                    in_toc = False
                    body_started = True
                else:
                    continue
            if is_toc_line(text) and not body_started:
                continue
            if is_toc_line(text) and re.search(r"\t\d+\s*$|\.{4,}\s*\d+\s*$", text):
                continue

            prev_section = current
            current = detect_section(text, current)
            if current != "cover" and current != prev_section:
                body_started = True

            style = (p.style.name if p.style else "") or ""
            # skip pure section titles echoed as body
            is_heading = any(text.startswith(n) for _, needles in SECTION_MATCHERS for n in needles)
            if is_heading:
                prev_text = text
                buckets[current]["prev_caption"] = text
                continue
            if style.lower().startswith("list"):
                buckets[current]["bullets"].append(text)
            else:
                buckets[current]["texts"].append(text)
            prev_text = text
            buckets[current]["prev_caption"] = text

        elif child.tag == qn("w:tbl"):
            if in_toc and not body_started:
                continue
            table = Table(child, doc)
            embeds = blip_embeds(child)
            if embeds and (body_started or current != "cover"):
                for rid in embeds:
                    images_by_section[current].append(rid)

            rows = table_matrix(table)
            if is_banner_table(rows):
                continue

            # Contenu dans un tableau imbriqué (ex. 1.4, géométrie, CCTP…)
            nested = nested_tables(table)
            if nested:
                label = ""
                caption_right = ""
                if len(rows) == 1 and rows[0]:
                    label = (rows[0][0] or "").strip()
                    if len(rows[0]) > 1:
                        caption_right = (rows[0][1] or "").strip()
                if label and ("document" in label.lower() or label.startswith("1.4")):
                    current = detect_section(
                        label if label.startswith("1.4") else "1.4. Documents de référence",
                        current,
                    )
                imported_nested = False
                for nt in nested:
                    nrows = table_matrix(nt)
                    if not nrows or is_banner_table(nrows):
                        continue
                    cap = label if label and "document" in label.lower() else ""
                    buckets[current]["tables"].append(free_table_from_rows(nrows, caption=cap))
                    imported_nested = True
                    body_started = True
                if imported_nested:
                    # Légende du bandeau (cellule droite) → texte d'intro si présente
                    if caption_right and not is_toc_line(caption_right):
                        buckets[current]["texts"].append(
                            f"{label}\n{caption_right}" if label else caption_right
                        )
                    elif label and current == "s1_4":
                        pass  # titre déjà porté par la section + tableau
                    continue

            if is_data_table(rows):
                caption = prev_text if prev_text and not any(
                    prev_text.startswith(n) for _, needles in SECTION_MATCHERS for n in needles
                ) else ""
                if len(caption) > 120:
                    caption = ""
                buckets[current]["tables"].append(free_table_from_rows(rows, caption=caption))
                body_started = True
                continue

            # 1x1 text box → narrative
            if len(rows) == 1 and len(rows[0]) == 1 and rows[0][0]:
                t = rows[0][0]
                if not is_toc_line(t):
                    buckets[current]["texts"].append(t)
                    prev_text = t[:100]
                    body_started = True
                continue

            # 1x2 label + contenu → toujours en texte (évite tableaux artificiels 1 ligne)
            if len(rows) == 1 and len(rows[0]) == 2:
                k, v = (rows[0][0] or "").strip(), (rows[0][1] or "").strip()
                if k and v:
                    if any(k.lower() == n.lower() or k.startswith(n) for _, needles in SECTION_MATCHERS for n in needles):
                        buckets[current]["texts"].append(v)
                    else:
                        buckets[current]["texts"].append(f"{k}\n{v}")
                    prev_text = (v or k)[:100]
                    body_started = True
                elif v:
                    buckets[current]["texts"].append(v)
                    body_started = True
                elif k:
                    buckets[current]["texts"].append(k)
                    body_started = True
                continue

            # grilles 2x2+ (ex. annexe 2)
            nr = len(rows)
            nc = max((len(r) for r in rows), default=0)
            if nr >= 2 and nc >= 2:
                buckets[current]["tables"].append(
                    free_table_from_rows(rows, caption=prev_text[:80] if prev_text else "")
                )
                # aussi une ligne texte récap si cellules courtes
                labels = [c for row in rows for c in row if c]
                if labels and all(len(c) < 80 for c in labels):
                    buckets[current]["texts"].append(" · ".join(labels))
                body_started = True

    return buckets, images_by_section, image_parts


def extract_cover_meta(docx_path: Path) -> dict:
    # Lightweight: reuse first paragraphs via zip XML for title
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    texts = []
    for para in re.split(r"</w:p>", xml)[:80]:
        parts = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", para)
        line = "".join(parts).strip()
        if line:
            texts.append(line)
    titre = next((t for t in texts if "Optimisation" in t), "Avis technique")
    lieu = texts[0] if texts else ""
    return {
        "title": titre,
        "lieu": lieu,
        "type": "Avis technique",
        "author": "COSTA PEREIRA Marco",
        "document_date": "30/07/2026",
        "status": "V2",
        "destinataire": "Valentin MONTEIL / GUINTOLI",
    }


def build_contents(buckets: dict, images_by_section: dict, rid_to_doc_id: dict[str, int]) -> dict:
    contents: dict = {}
    cover_fields = extract_cover_meta(DOCX)
    contents["cover"] = {
        "meta": {
            "block_type": "meta_document",
            "content": {"fields": cover_fields},
        }
    }

    for node in OUTLINE:
        sid = node["id"]
        if sid == "cover":
            continue
        if node.get("parent_id") is None:
            contents.setdefault(sid, {})
            continue

        bucket = buckets.get(sid) or {"texts": [], "bullets": [], "tables": []}
        texts = [t for t in (bucket.get("texts") or []) if t and not is_toc_line(t)]
        bullets = [b for b in (bucket.get("bullets") or []) if b and not is_toc_line(b)]
        tables = list(bucket.get("tables") or [])

        # Drop repeated titles / echo headings
        skip_titles = {node["title"].lower(), node["title"].split(". ", 1)[-1].lower()}
        skip_titles |= {s.lower() for s in (node.get("match") or [])}
        texts = [t for t in texts if t.lower() not in skip_titles and not any(t.startswith(n) for n in (node.get("match") or []))]
        bullets = [b for b in bullets if b.lower() not in skip_titles]

        # Expand multi-line method / list blocks into bullets when useful
        expanded_bullets = list(bullets)
        cleaned_texts = []
        for t in texts:
            lines = [ln.strip() for ln in t.splitlines() if ln.strip()]
            if sid in {"s1_5", "s1_2", "s1_3"} and len(lines) >= 3 and all(len(ln) < 220 for ln in lines):
                # première ligne = titre court éventuel
                if lines[0].lower() in {"méthodologie", "objectifs recherchés", "contraintes du site"}:
                    expanded_bullets.extend(lines[1:])
                else:
                    expanded_bullets.extend(lines)
            else:
                cleaned_texts.append(t)
        texts = cleaned_texts
        bullets = expanded_bullets

        slot: dict = {}
        if sid in {"s1_2", "s1_3", "s1_5"} and (bullets or texts):
            short = [t for t in bullets if t]
            long = [t for t in texts if len(t) >= 80]
            leftovers = [t for t in texts if len(t) < 80]
            if leftovers and not short:
                short = leftovers
                long = [t for t in texts if t not in leftovers]
            slot["liste"] = {"block_type": "bullet_list", "content": {"items": short}}
            slot["texte"] = {"block_type": "rich_text", "content": {"text": "\n\n".join(long)}}
        else:
            slot["texte"] = {"block_type": "rich_text", "content": {"text": "\n\n".join(texts)}}
            slot["liste"] = {"block_type": "bullet_list", "content": {"items": bullets}}

        # Tables → free_table (multi)
        if tables or node.get("prefer_table"):
            # Deduplicate identical tables
            unique_tables = []
            seen = set()
            for tbl in tables:
                key = json.dumps(tbl, ensure_ascii=False, sort_keys=True)
                if key in seen:
                    continue
                seen.add(key)
                unique_tables.append(tbl)
            slot["tableau"] = {
                "block_type": "free_table",
                "content": {"tables": unique_tables} if unique_tables else {"tables": []},
            }

        # Figures
        if sid in SECTIONS_WITH_FIGURES or node.get("prefer_media"):
            cards = []
            seen_docs: set[int] = set()
            for rid in images_by_section.get(sid) or []:
                doc_id = rid_to_doc_id.get(rid)
                if not doc_id or doc_id in seen_docs:
                    continue
                seen_docs.add(doc_id)
                cards.append(
                    {
                        "document_id": doc_id,
                        "caption": "",
                        "order": len(cards),
                        "display": "full_width",
                    }
                )
            slot["figures"] = {"block_type": "media_cards", "content": {"cards": cards}}

        if node.get("prefer_calculs"):
            slot["fiches"] = {
                "block_type": "calcul_fiches",
                "content": {"calcul_ids": [], "auto_from_binding": True},
            }

        contents[sid] = slot

    # Copy hypotheses table into 4.4 if empty but present in a5 (Word places it in annex)
    a5_tables = (contents.get("a5") or {}).get("tableau", {}).get("content", {}).get("tables") or []
    s44_tables = (contents.get("s4_4") or {}).get("tableau", {}).get("content", {}).get("tables") or []
    if not s44_tables and a5_tables:
        hypo = next(
            (
                t
                for t in a5_tables
                if t.get("headers")
                and any("param" in str(h).lower() for h in t["headers"])
            ),
            None,
        )
        if hypo and contents.get("s4_4"):
            contents["s4_4"]["tableau"] = {
                "block_type": "free_table",
                "content": {"tables": [dict(hypo, caption="Hypothèses de trafic")]},
            }

    return contents


def register_images(
    *,
    demande_id: int,
    affaire_ref: str,
    images_by_section: dict[str, list[str]],
    image_parts: dict[str, tuple[str, bytes]],
) -> dict[str, int]:
    """Save unique images to storage + demande_documents. Returns rId -> document_id."""
    from app.core.database import connect_db, get_db_path
    from app.services.demande_document_storage_service import save_affaire_document

    # Prefer first section occurrence for labeling
    rid_section: dict[str, str] = {}
    for sid, rids in images_by_section.items():
        for rid in rids:
            rid_section.setdefault(rid, sid)

    rid_to_doc: dict[str, int] = {}
    now = datetime.now().isoformat(timespec="seconds")

    with connect_db(get_db_path()) as conn:
        # Remove previous auto-imported avis figures for this demande (idempotent re-run)
        old = conn.execute(
            """
            SELECT id, stored_path FROM demande_documents
            WHERE demande_id = ? AND comment LIKE 'avis-import:%'
            """,
            (demande_id,),
        ).fetchall()
        for row in old:
            conn.execute("DELETE FROM demande_documents WHERE id = ?", (row["id"],))

        for rid, section_id in rid_section.items():
            part = image_parts.get(rid)
            if not part:
                continue
            filename, blob = part
            ext = Path(filename).suffix.lower() or ".png"
            safe_name = f"avis_nt_fig_{section_id}_{rid}{ext}"
            saved = save_affaire_document(affaire_ref, blob, safe_name)
            stored_path = saved["stored_path"]
            label = f"Figure avis · {section_id}"
            cur = conn.execute(
                """
                INSERT INTO demande_documents (
                    demande_id, document_type, is_received, version, document_date,
                    comment, stored_path, uploaded_at, created_at, updated_at
                ) VALUES (?, ?, 1, ?, NULL, ?, ?, ?, ?, ?)
                """,
                (
                    demande_id,
                    label,
                    Path(stored_path).name,
                    f"avis-import:{rid}",
                    stored_path,
                    now,
                    now,
                    now,
                ),
            )
            rid_to_doc[rid] = int(cur.lastrowid)
        conn.commit()

    return rid_to_doc


def main() -> None:
    from app.models.avis_technique import AvisInstanceUpdateSchema, AvisTemplateUpsertSchema
    from app.repositories.avis_technique_repository import AvisTechniqueRepository

    if not DOCX.exists():
        raise SystemExit(f"DOCX missing: {DOCX}")

    definition = build_template_definition()
    seed_payload = {
        "code": "optimisation_chaussee_rst",
        "label": "Avis technique — optimisation structure de chaussée",
        "version": 3,
        "docx_style_path": "",
        "definition": definition,
    }
    SEED.write_text(json.dumps(seed_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print("seed written", SEED)

    buckets, images_by_section, image_parts = walk_document(DOCX)
    print(
        "sections with tables:",
        {k: len(v["tables"]) for k, v in buckets.items() if v["tables"]},
    )
    print(
        "images by section:",
        {k: len(v) for k, v in images_by_section.items() if v},
    )

    repo = AvisTechniqueRepository()
    tpl = repo.upsert_template(
        AvisTemplateUpsertSchema(
            **{k: seed_payload[k] for k in ("code", "label", "version", "definition", "docx_style_path")}
        )
    )
    print("template", tpl["id"], "sections", len(tpl["definition"]["sections"]), "v", tpl.get("version"))

    inst = repo.get_instance(INSTANCE_ID, with_template=False)
    if not inst:
        raise SystemExit(f"instance {INSTANCE_ID} missing")
    demande_id = int(inst["demande_id"])

    rid_to_doc = register_images(
        demande_id=demande_id,
        affaire_ref=AFFAIRE_REF,
        images_by_section=images_by_section,
        image_parts=image_parts,
    )
    print("registered images", len(rid_to_doc))

    contents = build_contents(buckets, images_by_section, rid_to_doc)

    # Préserver / rebrancher Annexe 6 → série CAM1
    from app.services.avis_technique_bindings import resolve_calculs

    cam1 = resolve_calculs(demande_id, {"type_calcul": "alize", "nom_contains": "CAM1"})
    cam1_ids = [c["id"] for c in sorted(cam1, key=lambda x: str(x.get("nom_calcul") or ""))]
    a6 = dict(contents.get("a6") or {})
    a6["texte"] = {
        "block_type": "rich_text",
        "content": {
            "text": (
                "Les résultats sont présentés sous forme de fiches individuelles "
                "(méthode rationnelle LCPC-Sétra), une fiche par cas étudié."
            )
        },
    }
    a6["fiches"] = {
        "block_type": "calcul_fiches",
        "content": {"calcul_ids": cam1_ids, "auto_from_binding": True},
    }
    contents["a6"] = a6

    meta = dict(inst.get("meta") or {})
    cover_fields = contents.get("cover", {}).get("meta", {}).get("content", {}).get("fields", {})
    meta.update(cover_fields)
    meta["reference"] = inst["reference"]

    linked_ids = sorted({int(v) for v in rid_to_doc.values()})

    updated = repo.update_instance(
        INSTANCE_ID,
        AvisInstanceUpdateSchema(
            titre=cover_fields.get("title") or inst.get("titre"),
            meta=meta,
            contents=contents,
            statut="En rédaction",
            linked_document_ids=linked_ids,
        ),
    )
    print("instance", updated["reference"], "content keys", len(updated["contents"]))
    for sid in ("s1_1", "s1_4", "s1_5", "s2_2", "s2_5", "s3_1", "a2", "a3", "a5", "a6"):
        slot = updated["contents"].get(sid) or {}
        texte = (slot.get("texte") or {}).get("content", {}).get("text") or ""
        liste = (slot.get("liste") or {}).get("content", {}).get("items") or []
        tabs = (slot.get("tableau") or {}).get("content", {}).get("tables") or []
        figs = (slot.get("figures") or {}).get("content", {}).get("cards") or []
        fiches = (slot.get("fiches") or {}).get("content", {}).get("calcul_ids") or []
        print(
            f"{sid}: texte={len(texte)} liste={len(liste)} tables={len(tabs)} "
            f"figs={len(figs)} fiches={len(fiches)} | {texte[:70].replace(chr(10), ' / ')}"
        )


if __name__ == "__main__":
    main()
